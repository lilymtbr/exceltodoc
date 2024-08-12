import openpyxl
from docx import Document
import os
import tkinter as tk
from tkinter import filedialog, messagebox
from datetime import datetime
from docx.shared import Pt, RGBColor


# Função para formatar números com duas casas decimais
def format_decimal(value):
    try:
        # Tenta converter e formatar o valor como decimal
        return format(float(value), '.2f')
    except (TypeError, ValueError):
        # Se a conversão falhar, retorna o valor original
        return value
    
# Função para gerar o documento Word
def generate_word():
    try:
        excel_path = excel_entry.get()
        word_template_path = word_entry.get()
        output_path = output_entry.get()

        if not os.path.exists(excel_path):
            messagebox.showerror("Erro", f"Arquivo Excel não encontrado em {excel_path}")
            return

        if not os.path.exists(word_template_path):
            messagebox.showerror("Erro", f"Arquivo Word não encontrado em {word_template_path}")
            return

        # Carregar a planilha do Excel usando openpyxl
        workbook = openpyxl.load_workbook(excel_path, data_only=True)
        sheet_descr = workbook['Descrição dos ambientes']
        sheet_capa = workbook['Capa']
        sheet_carga_teq = workbook['Carga térmica equip.']
        sheet_paredes_janelas = workbook['Carga termica Paredes-janelas']
        sheet_total = workbook['Total']

        # Obter os valores das células J3 e K3
        temperatura_interna = sheet_descr['J3'].value
        temperatura_externa = sheet_descr['K3'].value

        # Obter os valores das células A40, A39, A38, A37, A36, A35 na aba Capa
        rev1 = sheet_capa['A40'].value
        rev2 = sheet_capa['A39'].value
        rev3 = sheet_capa['A38'].value
        rev4 = sheet_capa['A37'].value
        rev5 = sheet_capa['A36'].value
        rev6 = sheet_capa['A35'].value

        # Obter os valores das células B40, B39, B38, B37, B36, B35 na aba Capa
        data1 = sheet_capa['B40'].value
        data2 = sheet_capa['B39'].value
        data3 = sheet_capa['B38'].value
        data4 = sheet_capa['B37'].value
        data5 = sheet_capa['B36'].value
        data6 = sheet_capa['B35'].value

        # Obter os valores das células D40, D39, D38, D37, D36, D35 na aba Capa
        desc1 = sheet_capa['D40'].value
        desc2 = sheet_capa['D39'].value
        desc3 = sheet_capa['D38'].value
        desc4 = sheet_capa['D37'].value
        desc5 = sheet_capa['D36'].value
        desc6 = sheet_capa['D35'].value

         # Obter os valores das células G40, G39, G38, G37, G36, G35 na aba Capa
        feito1 = sheet_capa['G40'].value
        feito2 = sheet_capa['G39'].value
        feito3 = sheet_capa['G38'].value
        feito4 = sheet_capa['G37'].value
        feito5 = sheet_capa['G36'].value
        feito6 = sheet_capa['G35'].value

         # Obter os valores das células H40, H39, H38, H37, H36, H35 na aba Capa
        visto1 = sheet_capa['H40'].value
        visto2 = sheet_capa['H39'].value
        visto3 = sheet_capa['H38'].value
        visto4 = sheet_capa['H37'].value
        visto5 = sheet_capa['H36'].value
        visto6 = sheet_capa['H35'].value

          # Obter os valores das células I40, I39, I38, I37, I36, I35 na aba Capa
        aprov1 = sheet_capa['I40'].value
        aprov2 = sheet_capa['I39'].value
        aprov3 = sheet_capa['I38'].value
        aprov4 = sheet_capa['I37'].value
        aprov5 = sheet_capa['I36'].value
        aprov6 = sheet_capa['I35'].value

         # Obter os valores das células A49, D49, F49, G49, I49, A45, A51, H49 e I56 na aba Capa
        proj = sheet_capa['A49'].value
        verif = sheet_capa['D49'].value
        visto = sheet_capa['F49'].value
        aprov = sheet_capa['G49'].value
        data = sheet_capa['I49'].value
        empreendimento = sheet_capa['A45'].value
        subestacao = sheet_capa['A51'].value
        num_cliente = sheet_capa['A56'].value
        rev = sheet_capa['I56'].value

        # Obter os valores das células A3 a A20
        ref1 = sheet_capa['A3'].value
        ref2 = sheet_capa['A4'].value
        ref3 = sheet_capa['A5'].value
        ref4 = sheet_capa['A6'].value
        ref5 = sheet_capa['A7'].value
        ref6 = sheet_capa['A8'].value
        ref7 = sheet_capa['A9'].value
        ref8 = sheet_capa['A10'].value
        ref9 = sheet_capa['A11'].value
        ref10 = sheet_capa['A12'].value
        ref11 = sheet_capa['A13'].value
        ref12 = sheet_capa['A14'].value
        ref13 = sheet_capa['A15'].value
        ref14 = sheet_capa['A16'].value
        ref15 = sheet_capa['A17'].value
        ref16 = sheet_capa['A18'].value
        ref17 = sheet_capa['A19'].value
        ref18 = sheet_capa['A20'].value

        # Obter os valores das células C15 and C16
        carga_pnl_elet = sheet_carga_teq['C15'].value
        carga_quadro_elet = sheet_carga_teq['C16'].value

        # Obter os valores das células B11 e C11
        capac_bat_125vcc = sheet_carga_teq['B11'].value
        capac_bat_48vcc = sheet_carga_teq['C11'].value

        # Obter os valores das células especificadas
        ambiente1 = sheet_descr['B1'].value
        ambiente2 = sheet_descr['B9'].value
        ambiente3 = sheet_descr['B17'].value
        ambiente4 = sheet_descr['B25'].value
        ambiente5 = sheet_descr['B33'].value
        ambiente6 = sheet_descr['B41'].value

        # Obter os valores das células especificadas
        ambiente_br1 = sheet_descr['B1'].value
        ambiente_br2 = sheet_descr['B9'].value
        ambiente_br3 = sheet_descr['B17'].value
        ambiente_br4 = sheet_descr['B25'].value
        ambiente_br5 = sheet_descr['B33'].value
        ambiente_br6 = sheet_descr['B41'].value

        # Obter os valores das células especificadas
        area_n_values = [format_decimal(sheet_paredes_janelas[f'B{row}'].value) for row in range(4, 10)]
        area_s_values = [format_decimal(sheet_paredes_janelas[f'C{row}'].value) for row in range(4, 10)]
        area_e_values = [format_decimal(sheet_paredes_janelas[f'D{row}'].value) for row in range(4, 10)]
        area_o_values = [format_decimal(sheet_paredes_janelas[f'E{row}'].value) for row in range(4, 10)]

        carga_n_values = [format_decimal(sheet_paredes_janelas[f'B{row}'].value) for row in range(14, 20)]
        carga_s_values = [format_decimal(sheet_paredes_janelas[f'C{row}'].value) for row in range(14, 20)]
        carga_e_values = [format_decimal(sheet_paredes_janelas[f'D{row}'].value) for row in range(14, 20)]
        carga_o_values = [format_decimal(sheet_paredes_janelas[f'E{row}'].value) for row in range(14, 20)]
        carga_total_values = [format_decimal(sheet_paredes_janelas[f'F{row}'].value) for row in range(14, 20)]

        # Lista de células para os equipamentos e suas cargas
        equip_values = [sheet_carga_teq[f'B{row}'].value for row in range(15, 25)]  # B15 a B24 para equipamentos
        carga_eqto_values = [sheet_carga_teq[f'C{row}'].value for row in range(15, 25)]  # C15 a C24 para cargas dos equipamentos

        equip_data_placeholders = [
            '_q1a', '_q1b', '_q1c', '_q1d', '_q1e', '_q1f', '_q1g', '_q1h', '_q1i', '_q1j',
            '_q2a', '_q2b', '_q2c', '_q2d', '_q2e', '_q2f', '_q2g', '_q2h', '_q2i', '_q2j',
            '_q3a', '_q3b', '_q3c', '_q3d', '_q3e', '_q3f', '_q3g', '_q3h', '_q3i', '_q3j',
            '_q4a', '_q4b', '_q4c', '_q4d', '_q4e', '_q4f', '_q4g', '_q4h', '_q4i', '_q4j',
            '_q5a', '_q5b', '_q5c', '_q5d', '_q5e', '_q5f', '_q5g', '_q5h', '_q5i', '_q5j',
            '_q6a', '_q6b', '_q6c', '_q6d', '_q6e', '_q6f', '_q6g', '_q6h', '_q6i', '_q6j',
        ]

        # Lista de células correspondentes na aba "Carga térmica equip."
        equip_data_values = [
            'F17', 'G17', 'H17', 'I17', 'J17', 'K17', 'L17', 'M17', 'N17',
            'F18', 'G18', 'H18', 'I18', 'J18', 'K18', 'L18', 'M18', 'N18',
            'F19', 'G19', 'H19', 'I19', 'J19', 'K19', 'L19', 'M19', 'N19',
            'F20', 'G20', 'H20', 'I20', 'J20', 'K20', 'L20', 'M20', 'N20',
            'F21', 'G21', 'H21', 'I21', 'J21', 'K21', 'L21', 'M21', 'N21',
            'F22', 'G22', 'H22', 'I22', 'J22', 'K22', 'L22', 'M22', 'N22',
        ]

        equip_data_total_w_values = [sheet_carga_teq[f'O{row}'].value for row in range(17, 23)]  # B15 a B24 para equipamentos

        equip_data_total_carga_values = ['F25', 'G25', 'H25', 'I25', 'J25', 'K25']

        area_janela_n_values = [sheet_paredes_janelas[f'I{row}'].value for row in range(4, 10)]
        area_janela_s_values = [sheet_paredes_janelas[f'J{row}'].value for row in range(4, 10)]
        area_janela_e_values = [sheet_paredes_janelas[f'K{row}'].value for row in range(4, 10)]
        area_janela_o_values = [sheet_paredes_janelas[f'L{row}'].value for row in range(4, 10)]

        carga_janela_n_values = [format_decimal(sheet_paredes_janelas[f'B{row}'].value) for row in range(24, 30)]
        carga_janela_s_values = [format_decimal(sheet_paredes_janelas[f'C{row}'].value) for row in range(24, 30)]
        carga_janela_e_values = [format_decimal(sheet_paredes_janelas[f'D{row}'].value) for row in range(24, 30)]
        carga_janela_o_values = [format_decimal(sheet_paredes_janelas[f'E{row}'].value) for row in range(24, 30)]
        carga_janela_total_values = [format_decimal(sheet_paredes_janelas[f'F{row}'].value) for row in range(24, 30)]

        qde_pes_amb1 = ['_QtdePesMod1', '_QtdePesLento1', '_CalPes1']
        qde_pes_amb2 = ['_QtdePesMod2', '_QtdePesLento2', '_CalPes2']
        qde_pes_amb3 = ['_QtdePesMod3', '_QtdePesLento3', '_CalPes3']
        qde_pes_amb4 = ['_QtdePesMod4', '_QtdePesLento4', '_CalPes4']
        qde_pes_amb5 = ['_QtdePesMod5', '_QtdePesLento5', '_CalPes5']
        qde_pes_amb6 = ['_QtdePesMod6', '_QtdePesLento6', '_CalPes6']
        
        qte_pes_amb1 = [sheet_carga_teq[f'B{row}'].value for row in range(30, 33)]
        qte_pes_amb2 = [sheet_carga_teq[f'C{row}'].value for row in range(30, 33)]
        qte_pes_amb3 = [sheet_carga_teq[f'D{row}'].value for row in range(30, 33)]
        qte_pes_amb4 = [sheet_carga_teq[f'E{row}'].value for row in range(30, 33)]
        qte_pes_amb5 = [sheet_carga_teq[f'F{row}'].value for row in range(30, 33)]
        qte_pes_amb6 = [sheet_carga_teq[f'G{row}'].value for row in range(30, 33)]

        placeholders_ct_amb1 = ['_Amb_Par1', '_Amb_Pes1', '_Amb_Ilum1', '_Amb_Eq1', '_Amb_Jan1', '_Amb_Tot1']
        placeholders_ct_amb2 = ['_Amb_Par2', '_Amb_Pes2', '_Amb_Ilum2', '_Amb_Eq2', '_Amb_Jan2', '_Amb_Tot2']
        placeholders_ct_amb3 = ['_Amb_Par3', '_Amb_Pes3', '_Amb_Ilum3', '_Amb_Eq3', '_Amb_Jan3', '_Amb_Tot3']
        placeholders_ct_amb4 = ['_Amb_Par4', '_Amb_Pes4', '_Amb_Ilum4', '_Amb_Eq4', '_Amb_Jan4', '_Amb_Tot4']
        placeholders_ct_amb5 = ['_Amb_Par5', '_Amb_Pes5', '_Amb_Ilum5', '_Amb_Eq5', '_Amb_Jan5', '_Amb_Tot5']
        placeholders_ct_amb6 = ['_Amb_Par6', '_Amb_Pes6', '_Amb_Ilum6', '_Amb_Eq6', '_Amb_Jan6', '_Amb_Tot6']
        
        values_ct_amb1 = [format_decimal(sheet_total[f'B{row}'].value) for row in range(3, 9)]
        values_ct_amb2 = [format_decimal(sheet_total[f'C{row}'].value) for row in range(3, 9)]
        values_ct_amb3 = [format_decimal(sheet_total[f'D{row}'].value) for row in range(3, 9)]
        values_ct_amb4 = [format_decimal(sheet_total[f'E{row}'].value) for row in range(3, 9)]
        values_ct_amb5 = [format_decimal(sheet_total[f'F{row}'].value) for row in range(3, 9)]
        values_ct_amb6 = [format_decimal(sheet_total[f'G{row}'].value) for row in range(3, 9)]

        final_data_placeholders = [
            '_Amb_Kcal1', '_Amb_Kcal2', '_Amb_Kcal3', '_Amb_Kcal4', '_Amb_Kcal5', '_Amb_Kcal6',
            '_Amb_TR1', '_Amb_TR2', '_Amb_TR3', '_Amb_TR4', '_Amb_TR5', '_Amb_TR6',
            '_Amb_Btu1', '_Amb_Btu2', '_Amb_Btu3', '_Amb_Btu4', '_Amb_Btu5', '_Amb_Btu6',
        ]

        final_data_values = [
            'B12', 'C12', 'D12', 'E12', 'F12', 'G12',
            'B17', 'C17', 'D17', 'E17', 'F17', 'G17',
            'B18', 'C18', 'D18', 'E18', 'F18', 'G18',
        ]
        
        apar_values = ['B27', 'C27', 'D27', 'E27', 'F27', 'G27']
        apar_qtes = ['B28', 'C28', 'D28', 'E28', 'F28', 'G28']

        # Função para formatar datas
        def format_date(date):
            if isinstance(date, datetime):
                return date.strftime('%d/%m/%Y')
            return str(date)

        # Converter os valores para string, se necessário, substituindo None por string vazia
        temperatura_interna = str(temperatura_interna) if temperatura_interna is not None else ""
        temperatura_externa = str(temperatura_externa) if temperatura_externa is not None else ""
        rev1 = str(rev1) if rev1 is not None else ""
        rev2 = str(rev2) if rev2 is not None else ""
        rev3 = str(rev3) if rev3 is not None else ""
        rev4 = str(rev4) if rev4 is not None else ""
        rev5 = str(rev5) if rev5 is not None else ""
        rev6 = str(rev6) if rev6 is not None else ""
        data1 = format_date(data1) if data1 is not None else ""
        data2 = format_date(data2) if data2 is not None else ""
        data3 = format_date(data3) if data3 is not None else ""
        data4 = format_date(data4) if data4 is not None else ""
        data5 = format_date(data5) if data5 is not None else ""
        data6 = format_date(data6) if data6 is not None else ""
        desc1 = str(desc1) if desc1 is not None else ""
        desc2 = str(desc2) if desc2 is not None else ""
        desc3 = str(desc3) if desc3 is not None else ""
        desc4 = str(desc4) if desc4 is not None else ""
        desc5 = str(desc5) if desc5 is not None else ""
        desc6 = str(desc6) if desc6 is not None else ""
        feito1 = str(feito1) if feito1 is not None else ""
        feito2 = str(feito2) if feito2 is not None else ""
        feito3 = str(feito3) if feito3 is not None else ""
        feito4 = str(feito4) if feito4 is not None else ""
        feito5 = str(feito5) if feito5 is not None else ""
        feito6 = str(feito6) if feito6 is not None else ""
        visto1 = str(visto1) if visto1 is not None else ""
        visto2 = str(visto2) if visto2 is not None else ""
        visto3 = str(visto3) if visto3 is not None else ""
        visto4 = str(visto4) if visto4 is not None else ""
        visto5 = str(visto5) if visto5 is not None else ""
        visto6 = str(visto6) if visto6 is not None else ""
        aprov1 = str(aprov1) if aprov1 is not None else ""
        aprov2 = str(aprov2) if aprov2 is not None else ""
        aprov3 = str(aprov3) if aprov3 is not None else ""
        aprov4 = str(aprov4) if aprov4 is not None else ""
        aprov5 = str(aprov5) if aprov5 is not None else ""
        aprov6 = str(aprov6) if aprov6 is not None else ""
        proj = str(proj) if proj is not None else ""
        verif = str(verif) if verif is not None else ""
        visto = str(visto) if visto is not None else ""
        aprov = str(aprov) if aprov is not None else ""
        data = format_date(data) if data is not None else ""
        empreendimento = str(empreendimento) if empreendimento is not None else ""
        subestacao = str(subestacao) if subestacao is not None else ""
        num_cliente = str(num_cliente) if num_cliente is not None else ""
        rev = str(rev) if rev is not None else ""
        ref1 = str(ref1) if ref1 is not None else ""
        ref2 = str(ref2) if ref2 is not None else ""
        ref3 = str(ref3) if ref3 is not None else ""
        ref4 = str(ref4) if ref4 is not None else ""
        ref5 = str(ref5) if ref5 is not None else ""
        ref6 = str(ref6) if ref6 is not None else ""
        ref7 = str(ref7) if ref7 is not None else ""
        ref8 = str(ref8) if ref8 is not None else ""
        ref9 = str(ref9) if ref9 is not None else ""
        ref10 = str(ref10) if ref10 is not None else ""
        ref11 = str(ref11) if ref11 is not None else ""
        ref12 = str(ref12) if ref12 is not None else ""
        ref13 = str(ref13) if ref13 is not None else ""
        ref14 = str(ref14) if ref14 is not None else ""
        ref15 = str(ref15) if ref15 is not None else ""
        ref16 = str(ref16) if ref16 is not None else ""
        ref17 = str(ref17) if ref17 is not None else ""
        ref18 = str(ref18) if ref18 is not None else ""
        carga_pnl_elet = str(carga_pnl_elet) if carga_pnl_elet is not None else ""
        carga_quadro_elet = str(carga_quadro_elet) if carga_quadro_elet is not None else ""
        capac_bat_125vcc = str(capac_bat_125vcc) if capac_bat_125vcc is not None else ""
        capac_bat_48vcc = str(capac_bat_48vcc) if capac_bat_48vcc is not None else ""
        ambiente1 = str(ambiente1) if ambiente1 is not None else ""
        ambiente2 = str(ambiente2) if ambiente2 is not None else ""
        ambiente3 = str(ambiente3) if ambiente3 is not None else ""
        ambiente4 = str(ambiente4) if ambiente4 is not None else ""
        ambiente5 = str(ambiente5) if ambiente5 is not None else ""
        ambiente6 = str(ambiente6) if ambiente6 is not None else ""
        ambiente_br1 = str(ambiente_br1) if ambiente_br1 is not None else ""
        ambiente_br2 = str(ambiente_br2) if ambiente_br2 is not None else ""
        ambiente_br3 = str(ambiente_br3) if ambiente_br3 is not None else ""
        ambiente_br4 = str(ambiente_br4) if ambiente_br4 is not None else ""
        ambiente_br5 = str(ambiente_br5) if ambiente_br5 is not None else ""
        ambiente_br6 = str(ambiente_br6) if ambiente_br6 is not None else ""
        area_n_values = [str(value) if value is not None else "" for value in area_n_values]
        area_s_values = [str(value) if value is not None else "" for value in area_s_values]
        area_e_values = [str(value) if value is not None else "" for value in area_e_values]
        area_o_values = [str(value) if value is not None else "" for value in area_o_values]
        carga_n_values = [str(value) if value is not None else "" for value in carga_n_values]
        carga_s_values = [str(value) if value is not None else "" for value in carga_s_values]
        carga_e_values = [str(value) if value is not None else "" for value in carga_e_values]
        carga_o_values = [str(value) if value is not None else "" for value in carga_o_values]
        carga_total_values = [str(value) if value is not None else "" for value in carga_total_values]
        equip_values = [str(value) if value is not None else "" for value in equip_values]
        carga_eqto_values = [str(value) if value is not None else "" for value in carga_eqto_values]
        equip_data_values = [str(sheet_carga_teq[cell].value) if sheet_carga_teq[cell].value is not None else "" for cell in equip_data_values]
        equip_data_total_w_values = [str(value) if value is not None else "" for value in equip_data_total_w_values]
        equip_data_total_carga_values = [format_decimal(sheet_carga_teq[cell].value) if sheet_carga_teq[cell].value is not None else "" for cell in equip_data_total_carga_values]
        area_janela_n_values = [str(value) if value is not None else "" for value in area_janela_n_values]
        area_janela_s_values = [str(value) if value is not None else "" for value in area_janela_s_values]
        area_janela_e_values = [str(value) if value is not None else "" for value in area_janela_e_values]
        area_janela_o_values = [str(value) if value is not None else "" for value in area_janela_o_values]
        carga_janela_n_values = [str(value) if value is not None else "" for value in carga_janela_n_values]
        carga_janela_s_values = [str(value) if value is not None else "" for value in carga_janela_s_values]
        carga_janela_e_values = [str(value) if value is not None else "" for value in carga_janela_e_values]
        carga_janela_o_values = [str(value) if value is not None else "" for value in carga_janela_o_values]
        carga_janela_total_values = [str(value) if value is not None else "" for value in carga_janela_total_values]
        qte_pes_amb1 = [str(value) if value is not None else "" for value in qte_pes_amb1]
        qte_pes_amb2 = [str(value) if value is not None else "" for value in qte_pes_amb2]
        qte_pes_amb3 = [str(value) if value is not None else "" for value in qte_pes_amb3]
        qte_pes_amb4 = [str(value) if value is not None else "" for value in qte_pes_amb4]
        qte_pes_amb5 = [str(value) if value is not None else "" for value in qte_pes_amb5]
        qte_pes_amb6 = [str(value) if value is not None else "" for value in qte_pes_amb6]
        values_ct_amb1 = [str(value) if value is not None else "" for value in values_ct_amb1]
        values_ct_amb2 = [str(value) if value is not None else "" for value in values_ct_amb2]
        values_ct_amb3 = [str(value) if value is not None else "" for value in values_ct_amb3]
        values_ct_amb4 = [str(value) if value is not None else "" for value in values_ct_amb4]
        values_ct_amb5 = [str(value) if value is not None else "" for value in values_ct_amb5]
        values_ct_amb6 = [str(value) if value is not None else "" for value in values_ct_amb6]
        final_data_values = [format_decimal(sheet_total[cell].value) if sheet_total[cell].value is not None else "" for cell in final_data_values]
        apar_values = [str(sheet_total[cell].value) if sheet_total[cell].value is not None else "" for cell in apar_values]
        apar_qtes = [str(sheet_total[cell].value) if sheet_total[cell].value is not None else "" for cell in apar_qtes]

        # Carregar o documento Word
        document = Document(word_template_path)

        # Função para substituir texto no documento
        def replace_text_in_paragraph(paragraph, search_text, replace_text, font_size=None, font_color=None):
            full_text = ''.join(run.text for run in paragraph.runs)
            if search_text in full_text:
                # Limpa o parágrafo antes de inserir o novo texto
                for run in paragraph.runs:
                    run.text = ''
                # Insere o novo texto em um único run para manter a uniformidade
                new_run = paragraph.add_run(full_text.replace(search_text, replace_text))
                if font_size:
                    new_run.font.size = Pt(font_size)
                if font_color:
                    new_run.font.color.rgb = RGBColor(font_color[0], font_color[1], font_color[2])

        def replace_text_in_document(doc, search_text, replace_text, font_size=None, font_color=None):
            for paragraph in doc.paragraphs:
                replace_text_in_paragraph(paragraph, search_text, replace_text, font_size, font_color)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, search_text, replace_text, font_size, font_color)

        # Definir os valores a serem substituídos
        replace_text_in_document(document, '_TempInt', temperatura_interna, 12)
        replace_text_in_document(document, '_TempExt', temperatura_externa, 12)
        replace_text_in_document(document, '_1aRev', rev1)
        replace_text_in_document(document, '_2aRev', rev2)
        replace_text_in_document(document, '_3aRev', rev3)
        replace_text_in_document(document, '_4aRev', rev4)
        replace_text_in_document(document, '_5aRev', rev5)
        replace_text_in_document(document, '_6aRev', rev6)
        replace_text_in_document(document, '_1aData', data1)
        replace_text_in_document(document, '_2aData', data2)
        replace_text_in_document(document, '_3adata', data3)
        replace_text_in_document(document, '_4aData', data4)
        replace_text_in_document(document, '_5aData', data5)
        replace_text_in_document(document, '_6aData', data6)
        replace_text_in_document(document, '_Descrição1aRev', desc1)
        replace_text_in_document(document, '_Descrição2aRev', desc2)
        replace_text_in_document(document, '_Descrição3aRev', desc3)
        replace_text_in_document(document, '_Descrição4aRev', desc4)
        replace_text_in_document(document, '_Descrição5aRev', desc5)
        replace_text_in_document(document, '_Descrição6aRev', desc6)
        replace_text_in_document(document, '_1oFeito', feito1)
        replace_text_in_document(document, '_2oFeito', feito2)
        replace_text_in_document(document, '_3oFeito', feito3)
        replace_text_in_document(document, '_4oFeito', feito4)
        replace_text_in_document(document, '_5oFeito', feito5)
        replace_text_in_document(document, '_6oFeito', feito6)
        replace_text_in_document(document, '_1oVisto', visto1)
        replace_text_in_document(document, '_2oVisto', visto2)
        replace_text_in_document(document, '_3oVisto', visto3)
        replace_text_in_document(document, '_4oVisto', visto4)
        replace_text_in_document(document, '_5oVisto', visto5)
        replace_text_in_document(document, '_6oVisto', visto6)
        replace_text_in_document(document, '_1aAprov', aprov1)
        replace_text_in_document(document, '_2aAprov', aprov2)
        replace_text_in_document(document, '_3aAprov', aprov3)
        replace_text_in_document(document, '_4aAprov', aprov4)
        replace_text_in_document(document, '_5aAprov', aprov5)
        replace_text_in_document(document, '_6aAprov', aprov6)
        replace_text_in_document(document, '_Proj', proj)
        replace_text_in_document(document, '_Verif', verif)
        replace_text_in_document(document, '_Visto', visto)
        replace_text_in_document(document, '_Aprov', aprov)
        replace_text_in_document(document, '_Data', data)
        replace_text_in_document(document, '_Empreendimento', empreendimento)
        replace_text_in_document(document, '_Subestação', subestacao)
        replace_text_in_document(document, '_NumCliente', num_cliente)
        replace_text_in_document(document, '_Rev', rev)
        replace_text_in_document(document, '_Título', subestacao)
        replace_text_in_document(document, '1_REF', ref1, 12)
        replace_text_in_document(document, '_REF2', ref2, 12)
        replace_text_in_document(document, '_REF3', ref3, 12)
        replace_text_in_document(document, '_REF4', ref4, 12)
        replace_text_in_document(document, '_REF5', ref5, 12)
        replace_text_in_document(document, '_REF6', ref6, 12)
        replace_text_in_document(document, '_REF7', ref7, 12)
        replace_text_in_document(document, '_REF8', ref8, 12)
        replace_text_in_document(document, '_REF9', ref9, 12)
        replace_text_in_document(document, '_REF10', ref10, 12)
        replace_text_in_document(document, '_REF11', ref11, 12)
        replace_text_in_document(document, '_REF12', ref12, 12)
        replace_text_in_document(document, '_REF13', ref13, 12)
        replace_text_in_document(document, '_REF14', ref14, 12)
        replace_text_in_document(document, '_REF15', ref15, 12)
        replace_text_in_document(document, '_REF16', ref16, 12)
        replace_text_in_document(document, '_REF17', ref17, 12)
        replace_text_in_document(document, '_REF18', ref18, 12)
        replace_text_in_document(document, '_CargaPnlElet', carga_pnl_elet, 12)
        replace_text_in_document(document, '_CargaQuadroElet', carga_quadro_elet, 12)
        replace_text_in_document(document, '_CapacBat125Vcc', capac_bat_125vcc, 12)
        replace_text_in_document(document, '_CapacBat48Vcc', capac_bat_48vcc, 11)
        replace_text_in_document(document, '_Amb1', ambiente1, 11)
        replace_text_in_document(document, '_Amb2', ambiente2, 11)
        replace_text_in_document(document, '_Amb3', ambiente3, 11)
        replace_text_in_document(document, '_Amb4', ambiente4, 11)
        replace_text_in_document(document, '_Amb5', ambiente5, 11)
        replace_text_in_document(document, '_Amb6', ambiente6, 11)
        replace_text_in_document(document, '_Amb_White_1', ambiente_br1, 11, (255, 255, 255))
        replace_text_in_document(document, '_Amb_White_2', ambiente_br2, 11, (255, 255, 255))
        replace_text_in_document(document, '_Amb_White_3', ambiente_br3, 11, (255, 255, 255))
        replace_text_in_document(document, '_Amb_White_4', ambiente_br4, 11, (255, 255, 255))
        replace_text_in_document(document, '_Amb_White_5', ambiente_br5, 11, (255, 255, 255))
        replace_text_in_document(document, '_Amb_White_6', ambiente_br6, 11, (255, 255, 255))
        for i in range(6):
            replace_text_in_document(document, f'_ArParN{i+1}', area_n_values[i], 11)
            replace_text_in_document(document, f'_ArParS{i+1}', area_s_values[i], 11)
            replace_text_in_document(document, f'_ArParE{i+1}', area_e_values[i], 11)
            replace_text_in_document(document, f'_ArParO{i+1}', area_o_values[i], 11)
        for i in range(6):
            replace_text_in_document(document, f'_CarParN{i+1}', carga_n_values[i], 11)
            replace_text_in_document(document, f'_CarParS{i+1}', carga_s_values[i], 11)
            replace_text_in_document(document, f'_CarParE{i+1}', carga_e_values[i], 11)
            replace_text_in_document(document, f'_CarParO{i+1}', carga_o_values[i], 11)
            replace_text_in_document(document, f'_CarTotPar{i+1}', carga_total_values[i], 11)

        # Definir os placeholders para equipamentos e cargas
        equip_placeholders = ['_Equip_a', '_Equip_b', '_Equip_c', '_Equip_d', '_Equip_e', '_Equip_f', '_Equip_g', '_Equip_h', '_Equip_i', '_Equip_j']
        carga_eqto_placeholders = ['_CargaEqto_a', '_CargaEqto_b', '_CargaEqto_c', '_CargaEqto_d', '_CargaEqto_e', '_CargaEqto_f', '_CargaEqto_g', '_CargaEqto_h', '_CargaEqto_i', '_CargaEqto_j']

        # Executar a substituição para cada placeholder e valor correspondente
        for placeholder, value in zip(equip_placeholders, equip_values):
            replace_text_in_document(document, placeholder, value, 11)

        for placeholder, value in zip(carga_eqto_placeholders, carga_eqto_values):
            replace_text_in_document(document, placeholder, value, 11)

        for placeholder, value in zip(equip_data_placeholders, equip_data_values):
            replace_text_in_document(document, placeholder, value)

        for i in range(6):
            replace_text_in_document(document, f'_CWEqto{i+1}', equip_data_total_w_values[i], 11)

        for i in range(6):
            replace_text_in_document(document, f'_CalEqto{i+1}', equip_data_total_carga_values[i], 11)

        for i in range(6):
            replace_text_in_document(document, f'_ArJanN{i+1}', area_janela_n_values[i], 11)
            replace_text_in_document(document, f'_ArJanS{i+1}', area_janela_s_values[i], 11)
            replace_text_in_document(document, f'_ArJanE{i+1}', area_janela_e_values[i], 11)
            replace_text_in_document(document, f'_ArJanO{i+1}', area_janela_o_values[i], 11)

        for i in range(6):
            replace_text_in_document(document, f'_CarJanN{i+1}', carga_janela_n_values[i], 11)
            replace_text_in_document(document, f'_CarJanS{i+1}', carga_janela_s_values[i], 11)
            replace_text_in_document(document, f'_CarJanE{i+1}', carga_janela_e_values[i], 11)
            replace_text_in_document(document, f'_CarJanO{i+1}', carga_janela_o_values[i], 11)
            replace_text_in_document(document, f'_CalJAn{i+1}', carga_janela_total_values[i], 11)

        for placeholder, value in zip(qde_pes_amb1, qte_pes_amb1):
            replace_text_in_document(document, placeholder, value)
        for placeholder, value in zip(qde_pes_amb2, qte_pes_amb2):
            replace_text_in_document(document, placeholder, value)
        for placeholder, value in zip(qde_pes_amb3, qte_pes_amb3):
            replace_text_in_document(document, placeholder, value)
        for placeholder, value in zip(qde_pes_amb4, qte_pes_amb4):
            replace_text_in_document(document, placeholder, value)
        for placeholder, value in zip(qde_pes_amb5, qte_pes_amb5):
            replace_text_in_document(document, placeholder, value)
        for placeholder, value in zip(qde_pes_amb6, qte_pes_amb6):
            replace_text_in_document(document, placeholder, value)

        for placeholder, value in zip(placeholders_ct_amb1, values_ct_amb1):
            replace_text_in_document(document, placeholder, value)
        for placeholder, value in zip(placeholders_ct_amb2, values_ct_amb2):
            replace_text_in_document(document, placeholder, value)
        for placeholder, value in zip(placeholders_ct_amb3, values_ct_amb3):
            replace_text_in_document(document, placeholder, value)
        for placeholder, value in zip(placeholders_ct_amb4, values_ct_amb4):
            replace_text_in_document(document, placeholder, value)
        for placeholder, value in zip(placeholders_ct_amb5, values_ct_amb5):
            replace_text_in_document(document, placeholder, value)
        for placeholder, value in zip(placeholders_ct_amb6, values_ct_amb6):
            replace_text_in_document(document, placeholder, value)
        
        for placeholder, value in zip(final_data_placeholders, final_data_values):
            replace_text_in_document(document, placeholder, value, 11)

        for i in range(6):
            replace_text_in_document(document, f'_Aparelho{i+1}', apar_values[i], 11)
            replace_text_in_document(document, f'_Qtde{i+1}', apar_qtes[i], 11)

        # Salvar o documento Word preenchido
        document.save(output_path)

        messagebox.showinfo("Sucesso", f"Documento atualizado e salvo com sucesso em {output_path}")
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro: {e}")

# Função para selecionar o arquivo Excel
def select_excel_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    excel_entry.delete(0, tk.END)
    excel_entry.insert(0, file_path)

# Função para selecionar o arquivo Word
def select_word_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    word_entry.delete(0, tk.END)
    word_entry.insert(0, file_path)

# Função para selecionar o local para salvar o arquivo Word gerado
def select_output_file():
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    output_entry.delete(0, tk.END)
    output_entry.insert(0, file_path)

# Criar a interface gráfica
root = tk.Tk()
root.title("Gerador de Documentos")

tk.Label(root, text="Arquivo Excel:").grid(row=0, column=0, padx=10, pady=10, sticky="e")
excel_entry = tk.Entry(root, width=50)
excel_entry.grid(row=0, column=1, padx=10, pady=10)
tk.Button(root, text="Selecionar", command=select_excel_file).grid(row=0, column=2, padx=10, pady=10)

tk.Label(root, text="Arquivo Word (modelo):").grid(row=1, column=0, padx=10, pady=10, sticky="e")
word_entry = tk.Entry(root, width=50)
word_entry.grid(row=1, column=1, padx=10, pady=10)
tk.Button(root, text="Selecionar", command=select_word_file).grid(row=1, column=2, padx=10, pady=10)

tk.Label(root, text="Salvar documento como:").grid(row=2, column=0, padx=10, pady=10, sticky="e")
output_entry = tk.Entry(root, width=50)
output_entry.grid(row=2, column=1, padx=10, pady=10)
tk.Button(root, text="Selecionar", command=select_output_file).grid(row=2, column=2, padx=10, pady=10)

tk.Button(root, text="Gerar Documento", command=generate_word).grid(row=3, column=0, columnspan=3, pady=20)

root.mainloop()
